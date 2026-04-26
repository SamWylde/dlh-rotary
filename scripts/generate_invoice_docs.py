from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


INVOICE = {
    "title": "IN-KIND DONATION INVOICE",
    "subtitle": "Rotary Club of Downtown Lock Haven Website",
    "meta": [
        ("Invoice Number", "DLH-WEB-2026-04-23"),
        ("Invoice Date", "April 23, 2026"),
        ("Billing Period", "February 24, 2026 - April 22, 2026"),
        ("Prepared By", "Thomas Darby"),
        ("Provided To", "Rotary Club of Downtown Lock Haven"),
        ("Project", "Website design, development, content creation, content setup, launch support, and follow-up updates"),
        ("Website URL", "https://www.dlhrotary.org/"),
    ],
    "donation_note": [
        "This document is provided for record-keeping purposes as an in-kind donation of professional website services for the Rotary Club of Downtown Lock Haven.",
        "The hours and values below document the work completed for this project and the corresponding donated value. No payment is requested.",
    ],
    "services": [
        (
            "1",
            "Project planning, technical setup, development environment configuration, and deployment preparation",
            "5.0",
            "$150.00",
        ),
        (
            "2",
            "Website structure, content organization, page setup, site settings, and administrative configuration",
            "6.0",
            "$180.00",
        ),
        (
            "3",
            "Website design and front-end development, including homepage design, themed layouts, responsive page builds, reusable interface components, and content presentation",
            "10.0",
            "$300.00",
        ),
        (
            "4",
            "Interactive functionality, including forms, event features, RSVP workflow, navigation, metadata, and supporting integrations",
            "5.0",
            "$150.00",
        ),
        (
            "5",
            "Quality assurance, bug fixing, accessibility improvements, performance optimization, and launch polish",
            "4.0",
            "$120.00",
        ),
        (
            "6",
            "Content creation, integration, editing, and refinement based on supplied club materials, agendas, meeting notes, and follow-up edits",
            "2.0",
            "$60.00",
        ),
        (
            "7",
            "Post-launch updates and enhancements, including analytics setup and Flags of Honor campaign-related additions",
            "4.0",
            "$120.00",
        ),
        ("", "Total Project Hours", "36.0", ""),
        ("", "Value of Donated Services", "", "$1,080.00"),
        ("", "Donated Services Credit", "", "-$1,080.00"),
        ("", "Total Due", "", "$0.00"),
    ],
    "project_summary": [
        "Design and development of a custom Rotary Club website tailored to the club's current needs",
        "Content creation, editing, and organization based on supplied club materials, announcements, agendas, meeting notes, and campaign information",
        "Creation of website sections for pages, events, announcements, projects, documents, members, and club information",
        "Responsive layout and themed page styling for desktop and mobile use",
        "Setup of interactive features such as contact and interest forms, an event calendar, and RSVP functionality",
        "Search engine metadata, sitemap support, analytics setup, and deployment configuration",
        "Continued content and feature updates after launch, including support for the Flags of Honor campaign",
    ],
    "technical_summary": [
        "Approximately 11,300 lines of custom code and configuration",
        "More than 160 tracked project files supporting the website build",
        "22 routed pages and views",
        "8 major content sections and 3 site-wide settings areas",
        "68 reusable interface and feature components",
    ],
    "closing_note": "These services were donated in full. This invoice is intended only as a record of the in-kind contribution and no payment is required.",
}


ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "DLH-Rotary-Website-Invoice-2026-04-23.docx"
PDF_PATH = ROOT / "DLH-Rotary-Website-Invoice-2026-04-23.pdf"

ACCENT = RGBColor(26, 68, 120)
LIGHT_FILL = "EAF1F8"
GRID_COLOR = colors.HexColor("#C9D4E2")
HEADER_BG = colors.HexColor("#EAF1F8")


def set_cell_shading(cell, fill: str) -> None:
    cell_props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell_props.append(shading)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_doc_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(10.5)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles[f"Heading {level}"]
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.color.rgb = ACCENT


def add_title_block(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(INVOICE["title"])
    title_run.bold = True
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = ACCENT

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run(INVOICE["subtitle"])
    subtitle_run.font.name = "Aptos"
    subtitle_run.font.size = Pt(13)
    subtitle_run.italic = True


def add_meta_table(document: Document) -> None:
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(1.6)
    table.columns[1].width = Inches(4.9)

    for label, value in INVOICE["meta"]:
        row = table.add_row()
        row.cells[0].text = label
        if label == "Website URL":
            paragraph = row.cells[1].paragraphs[0]
            add_hyperlink(paragraph, value, value)
        else:
            row.cells[1].text = value
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.name = "Aptos"
        if row.cells[1].paragraphs[0].runs:
            row.cells[1].paragraphs[0].runs[0].font.name = "Aptos"

    document.add_paragraph()


def add_body_paragraphs(document: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_service_table(document: Document) -> None:
    headers = ["Item", "Service Description", "Project Hours", "Donated Value"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    widths = [0.5, 4.6, 1.0, 1.1]
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        header_cells[idx].paragraphs[0].runs[0].bold = True
        header_cells[idx].paragraphs[0].runs[0].font.name = "Aptos"
        set_cell_shading(header_cells[idx], LIGHT_FILL)

    for row_data in INVOICE["services"]:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            row.cells[idx].text = value
            if row.cells[idx].paragraphs[0].runs:
                row.cells[idx].paragraphs[0].runs[0].font.name = "Aptos"
        for idx in (2, 3):
            row.cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for row_index in range(len(INVOICE["services"]) - 3, len(INVOICE["services"]) + 1):
        table.rows[row_index].cells[1].paragraphs[0].runs[0].bold = True
        amount_cell = table.rows[row_index].cells[3]
        if amount_cell.paragraphs[0].runs:
            amount_cell.paragraphs[0].runs[0].bold = True

    document.add_paragraph()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.name = "Aptos"


def create_docx() -> None:
    document = Document()
    set_doc_defaults(document)
    add_title_block(document)
    add_meta_table(document)

    add_heading(document, "Donation Note")
    add_body_paragraphs(document, INVOICE["donation_note"])
    document.add_paragraph()

    add_heading(document, "Donated Services Summary")
    add_service_table(document)

    add_heading(document, "Project Summary")
    add_bullets(document, INVOICE["project_summary"])
    document.add_paragraph()

    add_heading(document, "Technical Summary")
    add_bullets(document, INVOICE["technical_summary"])
    document.add_paragraph()

    add_heading(document, "Closing Note")
    add_body_paragraphs(document, [INVOICE["closing_note"]])
    document.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "InvoiceTitle",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=28,
            textColor=colors.HexColor("#1A4478"),
            alignment=TA_LEFT,
            spaceAfter=6,
        ),
        "subtitle": ParagraphStyle(
            "InvoiceSubtitle",
            parent=styles["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=12,
            leading=16,
            spaceAfter=14,
        ),
        "heading": ParagraphStyle(
            "InvoiceHeading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#1A4478"),
            spaceBefore=10,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "InvoiceBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            spaceAfter=8,
        ),
        "bullet": ParagraphStyle(
            "InvoiceBullet",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            leftIndent=14,
            firstLineIndent=-8,
            bulletIndent=0,
            spaceAfter=4,
        ),
        "small": ParagraphStyle(
            "InvoiceSmall",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12,
        ),
    }


def create_pdf() -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=LETTER,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
    )

    story = [
        Paragraph(INVOICE["title"], styles["title"]),
        Paragraph(INVOICE["subtitle"], styles["subtitle"]),
    ]

    meta_rows = []
    for label, value in INVOICE["meta"]:
        if label == "Website URL":
            value_paragraph = Paragraph(f'<link href="{value}" color="blue">{value}</link>', styles["small"])
        else:
            value_paragraph = Paragraph(value, styles["small"])
        meta_rows.append([Paragraph(f"<b>{label}</b>", styles["small"]), value_paragraph])
    meta_table = Table(meta_rows, colWidths=[1.55 * inch, 5.0 * inch], hAlign="LEFT")
    meta_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.extend([meta_table, Spacer(1, 0.14 * inch)])

    story.append(Paragraph("Donation Note", styles["heading"]))
    for paragraph in INVOICE["donation_note"]:
        story.append(Paragraph(paragraph, styles["body"]))

    story.append(Paragraph("Donated Services Summary", styles["heading"]))
    service_rows = [
        [
            Paragraph("<b>Item</b>", styles["small"]),
            Paragraph("<b>Service Description</b>", styles["small"]),
            Paragraph("<b>Project Hours</b>", styles["small"]),
            Paragraph("<b>Donated Value</b>", styles["small"]),
        ]
    ]
    for row in INVOICE["services"]:
        service_rows.append([Paragraph(value or "&nbsp;", styles["small"]) for value in row])

    service_table = Table(
        service_rows,
        colWidths=[0.45 * inch, 4.35 * inch, 0.85 * inch, 1.1 * inch],
        hAlign="LEFT",
        repeatRows=1,
    )
    service_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), HEADER_BG),
                ("GRID", (0, 0), (-1, -1), 0.5, GRID_COLOR),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (2, 1), (3, -1), "RIGHT"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("FONTNAME", (0, 8), (-1, -1), "Helvetica-Bold"),
            ]
        )
    )
    story.extend([service_table, Spacer(1, 0.14 * inch)])

    story.append(Paragraph("Project Summary", styles["heading"]))
    for item in INVOICE["project_summary"]:
        story.append(Paragraph(f"- {item}", styles["bullet"]))

    story.append(Paragraph("Technical Summary", styles["heading"]))
    for item in INVOICE["technical_summary"]:
        story.append(Paragraph(f"- {item}", styles["bullet"]))

    story.append(Paragraph("Closing Note", styles["heading"]))
    story.append(Paragraph(INVOICE["closing_note"], styles["body"]))

    doc.build(story)


def main() -> None:
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    create_docx()
    create_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
